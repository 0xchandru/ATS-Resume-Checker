import re
import logging
import os
from typing import Any, List, Dict

logger = logging.getLogger(__name__)


def parse_resume(file_path: str) -> dict:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext == ".docx":
        return _parse_docx(file_path)
    elif ext == ".txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ─────────────────────────────────────────────────────────────────────────────
# Text Post-processing Utilities
# ─────────────────────────────────────────────────────────────────────────────

def _join_hyphenated_linebreaks(text: str) -> str:
    """
    Join words that were hyphenated across line breaks in PDFs.

    PDFs break long words at line boundaries with a soft hyphen:
      'Velocirap-\\ntor'          →  'Velociraptor'
      'Malware-\\nBazaar'         →  'MalwareBazaar'
      'system admin-\\nistration' →  'system administration'
      'case man-\\nagement'       →  'case management'
      'identifi-\\ncation'        →  'identification'

    We only join when the character immediately before the hyphen is a letter
    and the first character on the next line is also a letter — i.e., it is
    unambiguously a wrapped word, not a real compound hyphen followed by an
    independent sentence.
    """
    return re.sub(
        r'([A-Za-z])-[ \t]*\n[ \t]*([A-Za-z])',
        lambda m: m.group(1) + m.group(2),
        text,
    )


# ── Inline text artifact patterns (mirrors analyze._INLINE_FIXUPS) ────────────
# Applied to raw text in the parser so that rich_text HTML is also clean.
_TEXT_INLINE_FIXUPS: List[tuple] = [
    # email/URL domain directly followed by capitalised word — insert space:
    #   foo@bar.comPortfolio  →  foo@bar.com Portfolio
    (re.compile(
        r'(\.(com|net|org|io|me|uk|edu|co|app|dev|tech|ai|gov|mil|ca|au|in))([A-Z][a-z])'
    ), r'\1 \3'),
    # "Verify" stuck to another word (PDF cert badge artifact):
    #   VerifyGitHub → GitHub
    (re.compile(r'\bVerify(?=[A-Z][a-zA-Z])'), ''),
    # Common link labels stuck to adjacent words:
    #   GitHubMyProject → GitHub MyProject
    (re.compile(
        r'\b(GitHub|LinkedIn|Portfolio|TryHackMe|HackTheBox|LeetCode)(?=[A-Z][a-z])'
    ), r'\1 '),
    # Known social/professional domains stuck directly to a preceding word char
    # (the separator character was not captured during PDF extraction):
    #   0xchandrulinkedin.com  →  0xchandru linkedin.com
    #   chandraprakash.mePortfolio → handled by first rule
    # Uses lookbehind so we only insert space when NOT already preceded by a separator.
    (re.compile(
        r'(?<=[A-Za-z0-9])'
        r'(linkedin\.com|github\.com|github\.io|twitter\.com|x\.com|'
        r'medium\.com|youtube\.com|stackoverflow\.com|hashnode\.com|'
        r'dev\.to|tryhackme\.com|hackthebox\.com|leetcode\.com)',
        re.IGNORECASE,
    ), r' \1'),
]

# ── Section header detection (shared across normalization functions) ──────────
# Matches ALL-CAPS headings like "TECHNICAL SKILLS" or "PROJECTS & LAB EXPERIENCE"
_SECTION_HEADER_RE = re.compile(r'^[A-Z][A-Z0-9 &/|\-]{2,}$')
_BULLET_RE = re.compile(r'^[\u2022\u2013\u2014\-\*]\s|^\d+\.\s')
_SENTENCE_END_RE = re.compile(r'[.!?][\'")\]]*\s*$')

# Contact-info line detector: matches ONLY lines that contain an email address
# OR a profile-level URL (not repo-level URLs or tech-stack bullets).
# Profile URL = domain/1-segment (e.g. linkedin.com/in/handle, github.com/user)
# Repo URL   = domain/user/repo  — deliberately NOT matched.
_EMAIL_RE = re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9\-\.]+\.[a-z]{2,}')
_PROFILE_URL_RE = re.compile(
    r'linkedin\.com/(?:in/)[^/\s|·•]+'       # LinkedIn profile path
    r'|github\.com/[^/\s|·•]+(?:\s|[|·•]|$)' # GitHub handle (≤1 path segment)
    r'|twitter\.com/[^/\s|·•]+'
    r'|x\.com/[^/\s|·•]+',
    re.IGNORECASE,
)


def _deduplicate_contact_header(lines: List[str]) -> List[str]:
    """
    PDFs with a multi-row contact header produce 2-3 near-identical lines at
    the top.  Merge all contact-info lines found BEFORE the first section
    header into one canonical line.

    A "contact-info line" must have an email address OR a profile-level URL
    (github.com/handle, linkedin.com/in/handle).  This deliberately excludes:
      • repo URLs (github.com/user/reponame — 2 path segments)
      • tool-stack lines (Splunk • Wazuh • … — no URL or email)

    Example (3 → 1 line):
        chandraprakash.soc@gmail.com|linkedin.com/in/…|chandraprakash.me
        chandraprakash.soc@gmail.com|github.com/0xchandru linkedin.com/in/…|…
        chandraprakash.soc@gmail.com|github.com/0xchandru linkedin.com/in/…|…
    →
        chandraprakash.soc@gmail.com | linkedin.com/in/chandraprakash-soc | chandraprakash.me | github.com/0xchandru
    """
    # Stop scanning at the first ALL-CAPS section header (after line 2)
    header_end = min(30, len(lines))
    for i, ln in enumerate(lines[2:30], start=2):
        if _SECTION_HEADER_RE.match(ln.strip()) and len(ln.strip()) > 4:
            header_end = i
            break

    contact_indices = [
        i for i, ln in enumerate(lines[:header_end])
        if ln.strip() and (
            _EMAIL_RE.search(ln.strip())
            or _PROFILE_URL_RE.search(ln.strip())
        )
    ]
    if len(contact_indices) <= 1:
        return lines

    merged_tokens: List[str] = []
    seen_norm: set = set()
    for idx in contact_indices:
        line = lines[idx].strip()
        # Split on common separators: |, ·, •, or 2+ whitespace
        parts = re.split(r'[|·•]|\s{2,}', line)
        for part in parts:
            part = part.strip().strip('|·•, ')
            if not part:
                continue
            norm = re.sub(r'[\s/]+', '', part).lower()
            if norm and norm not in seen_norm:
                seen_norm.add(norm)
                merged_tokens.append(part)

    result = list(lines)
    result[contact_indices[0]] = ' | '.join(merged_tokens)
    for idx in contact_indices[1:]:
        result[idx] = None   # sentinel — filtered out after processing
    return result


def _merge_wrapped_paragraphs(lines: List[str]) -> List[str]:
    """
    PDF character extraction yields one *visual* line per text entry.  When a
    paragraph wraps across visual lines each chunk is a separate entry
    separated by blank lines.  This reconstructs the original paragraphs.

    Algorithm: multi-pass until no more merges are possible.  In each pass,
    whenever a line fragment (no sentence terminator) is separated from its
    continuation by exactly ONE blank line, the two are joined.  Cascade
    merges (A→B→C) are handled automatically because the merged A+B line is
    re-evaluated in the next pass.

    Merge triggers:
      • Previous line ends with comma   (list continuation)
      • Previous line ends with colon   (skill label + values)
      • Next line starts with lowercase  (mid-sentence word wrap)
      • Previous line has no terminal punctuation AND both lines are short
        (generic word-wrap heuristic — handles "…and MITRE\\n\\nATT&CK-…")

    Safe-guards (never merge across):
      • ALL-CAPS section headers
      • Bullet-point lines
      • Gaps of more than one blank line (intentional visual break)
    """
    _HDR = _SECTION_HEADER_RE
    _BUL = _BULLET_RE
    _END = _SENTENCE_END_RE

    # Filter out None sentinels from contact deduplication
    lines = [ln for ln in lines if ln is not None]

    changed = True
    while changed:
        changed = False
        new_lines: List[str] = []
        i = 0
        while i < len(lines):
            stripped = lines[i].strip()

            # Blank lines and un-mergeable lines pass through unchanged
            if not stripped or _HDR.match(stripped) or _BUL.match(stripped):
                new_lines.append(lines[i])
                i += 1
                continue

            # Find next non-empty line and count blanks in between
            j = i + 1
            blanks = 0
            while j < len(lines) and not lines[j].strip():
                blanks += 1
                j += 1

            # Too far away or no successor → don't merge
            if j >= len(lines) or blanks > 1:
                new_lines.append(lines[i])
                i += 1
                continue

            next_stripped = lines[j].strip()

            # Never merge into a header or bullet
            if _HDR.match(next_stripped) or _BUL.match(next_stripped):
                new_lines.append(lines[i])
                i += 1
                continue

            ends_sentence = bool(_END.search(stripped))
            ends_comma    = stripped.endswith(',')
            ends_colon    = stripped.endswith(':')
            next_lc       = bool(re.match(r'^[a-z]', next_stripped))
            looks_wrapped = (
                not ends_sentence
                and re.search(r'[A-Za-z0-9,]$', stripped)
                and len(stripped) < 100
                and len(next_stripped) < 100
            )

            should_merge = (
                ends_comma
                or ends_colon
                or (not ends_sentence and next_lc)
                or looks_wrapped
            )

            if should_merge:
                merged = stripped + ' ' + next_stripped
                new_lines.append(merged)
                # Preserve the blank lines (they collapse later)
                new_lines.extend([''] * blanks)
                # Skip lines[i] (consumed into merged) and lines[j] (merged source)
                i = j + 1
                changed = True
            else:
                new_lines.append(lines[i])
                i += 1

        lines = new_lines

    return lines


def _clean_extracted_text(text: str) -> str:
    """
    Apply inline artifact fixups to raw extracted text.
    Called on both pdfplumber and PyMuPDF raw text before storage.
    """
    for pat, repl in _TEXT_INLINE_FIXUPS:
        text = pat.sub(repl, text)
    # Remove lines that are only "Verify" repeated (cert badge lines)
    text = re.sub(r'^(Verify\s*)+$', '', text, flags=re.IGNORECASE | re.MULTILINE)
    # Collapse 3+ blank lines → 1
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _normalize_raw_text(text: str) -> str:
    """
    Full normalization pipeline applied to all parsed raw text AFTER the
    basic inline fixups.  Fixes document-level structural problems that
    _clean_extracted_text cannot address with simple regex passes.

    Pipeline:
      1. Basic inline fixups (URL concatenation, artifact chars)
      2. Contact-header deduplication
      3. Paragraph fragment reconstruction
      4. Excess blank-line collapsing (max 1 blank line between any two lines)
    """
    # Step 1: basic inline cleanup
    text = _clean_extracted_text(text)

    # Step 2–3 work on a line list
    lines = text.split('\n')
    lines = _deduplicate_contact_header(lines)
    lines = _merge_wrapped_paragraphs(lines)

    # Step 4: collapse consecutive blank lines to at most one
    text = '\n'.join(lines)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def _clean_html_text_nodes(html: str) -> str:
    """
    Apply inline artifact fixups to the TEXT NODES of an HTML string —
    i.e. only to content between tags, not to URLs or tag attributes.
    """
    parts = re.split(r'(<[^>]+>)', html)
    for i in range(0, len(parts), 2):  # only text nodes (even indices)
        node = parts[i]
        for pat, repl in _TEXT_INLINE_FIXUPS:
            node = pat.sub(repl, node)
        # Remove pure "Verify" runs within a text node
        node = re.sub(r'\bVerify(\s*Verify)*\b', '', node, flags=re.IGNORECASE)
        parts[i] = node
    return "".join(parts)


# Separator characters that commonly appear between links on the same line.
# Ordered from most specific (multi-char) to least so we match greedily.
_LINK_SEP_CHARS = ["—", "–", " | ", "|", " · ", "·", "•", "/", ","]


def _clean_link_label(text: str) -> str:
    """Strip separator characters and whitespace from both ends of a link label."""
    s = text.strip()
    changed = True
    while changed:
        changed = False
        for ch in ["—", "–", "|", "·", "•", "/", ","]:
            if s.startswith(ch) or s.endswith(ch):
                s = s.strip(ch).strip()
                changed = True
    return s


def _detect_separator_in_gap(page, rect_left, rect_right) -> str:
    """
    Read the text that sits *between* two adjacent link rectangles on the
    same line and return whichever separator character is found there.
    Falls back to '|' if nothing matches.
    """
    try:
        import fitz
        gap = fitz.Rect(
            rect_left[2],                          # left rect x1
            min(rect_left[1], rect_right[1]) - 2,  # top with small margin
            rect_right[0],                          # right rect x0
            max(rect_left[3], rect_right[3]) + 2,  # bottom with small margin
        )
        gap_text = page.get_text("text", clip=gap).strip()
        for sep in _LINK_SEP_CHARS:
            if sep in gap_text:
                return sep.strip() or "|"
    except Exception:
        pass
    return "|"


def _extract_pdf_links(file_path: str) -> List[Dict]:
    """
    Extract all hyperlinks from a PDF using PyMuPDF annotations.

    Each returned dict contains:
      text        – clean label (separator chars stripped)
      url         – destination URI
      page        – 1-based page number
      rect        – [x0, y0, x1, y1] of the clickable area
      line_group  – key identifying which visual line this link sits on
                    (format: "p{page}-y{rounded_y}")
      separator   – separator character detected between this link and
                    the next one on the same line ('' for the last link)
      position    – 0-based index within its line group
    """
    links: List[Dict] = []
    try:
        import fitz
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            page_links = page.get_links()
            if not page_links:
                continue

            words = page.get_text("words")  # (x0,y0,x1,y1,word,blk,ln,wi)

            raw_links: List[Dict] = []
            for link in page_links:
                if link.get("kind") != 2:   # LINK_URI = 2
                    continue
                uri = link.get("uri", "").strip()
                if not uri or uri.startswith("#"):
                    continue
                rect = link.get("from")
                if rect is None:
                    continue

                # Collect words whose bounding boxes overlap the link rect
                link_words = []
                for w in words:
                    wx0, wy0, wx1, wy1 = w[0], w[1], w[2], w[3]
                    if wx0 < rect.x1 and wx1 > rect.x0 and wy0 < rect.y1 and wy1 > rect.y0:
                        link_words.append(w[4])

                label = _clean_link_label(" ".join(link_words))
                rect_list = [rect.x0, rect.y0, rect.x1, rect.y1]
                y_center = (rect.y0 + rect.y1) / 2

                raw_links.append({
                    "text": label,
                    "url": uri,
                    "page": page_num + 1,
                    "rect": rect_list,
                    "_y_center": y_center,
                    "_page_obj": page,
                })

            # ── Group links that share the same visual line (within 4 px) ──
            raw_links.sort(key=lambda l: (l["_y_center"], l["rect"][0]))

            groups: List[List[Dict]] = []
            for lnk in raw_links:
                placed = False
                for grp in groups:
                    if abs(lnk["_y_center"] - grp[0]["_y_center"]) <= 4:
                        grp.append(lnk)
                        placed = True
                        break
                if not placed:
                    groups.append([lnk])

            for grp in groups:
                grp.sort(key=lambda l: l["rect"][0])   # left → right
                y_key = round(grp[0]["_y_center"])
                group_key = f"p{page_num + 1}-y{y_key}"

                for pos, lnk in enumerate(grp):
                    # Detect separator in the gap between this and the next link
                    sep = ""
                    if pos < len(grp) - 1:
                        sep = _detect_separator_in_gap(
                            lnk["_page_obj"], lnk["rect"], grp[pos + 1]["rect"]
                        )

                    links.append({
                        "text": lnk["text"],
                        "url": lnk["url"],
                        "page": lnk["page"],
                        "rect": lnk["rect"],
                        "line_group": group_key,
                        "separator": sep,
                        "position": pos,
                    })

        doc.close()
    except Exception as e:
        logger.warning("PDF link extraction failed: %s", e)
    return links


def _build_link_rect_index(links: List[Dict]) -> Dict[int, List]:
    """Index links by page number for fast lookup."""
    idx: Dict[int, List] = {}
    for lnk in links:
        pg = lnk["page"]
        idx.setdefault(pg, []).append(lnk)
    return idx


def _span_in_link(span_bbox, page_links: List[Dict]) -> str:
    """Return the URL of the link with the greatest overlap with this span's bounding box.

    Using maximum-intersection-area matching prevents adjacent links on the
    same line from all resolving to the first link's URL (the old first-match
    bug).  Each span is now assigned to whichever link rectangle it overlaps
    with the most.
    """
    if not page_links:
        return ""
    sx0, sy0, sx1, sy1 = span_bbox
    best_url = ""
    best_area = 0.0
    for lnk in page_links:
        rx0, ry0, rx1, ry1 = lnk["rect"]
        # Intersection rectangle
        ix0 = max(sx0, rx0)
        iy0 = max(sy0, ry0)
        ix1 = min(sx1, rx1)
        iy1 = min(sy1, ry1)
        if ix1 > ix0 and iy1 > iy0:
            area = (ix1 - ix0) * (iy1 - iy0)
            if area > best_area:
                best_area = area
                best_url = lnk["url"]
    return best_url


# ─────────────────────────────────────────────────────────────────────────────
# PDF Parser
# ─────────────────────────────────────────────────────────────────────────────

def _parse_pdf(file_path: str) -> dict:
    """Main PDF entry point: builds rich HTML + raw text with all enhancements."""

    # ── Extract links first (we'll use them in both HTML and raw_text) ──────
    links = _extract_pdf_links(file_path)
    link_index = _build_link_rect_index(links)

    rich_text = ""
    try:
        import fitz
        doc = fitz.open(file_path)
        html_parts = []

        # Calculate average font size across the document for heading detection
        sizes = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            if span.get("text", "").strip():
                                sizes.append(span.get("size", 12.0))
        avg_size = sum(sizes) / len(sizes) if sizes else 12.0

        for page_num, page in enumerate(doc):
            page_links = link_index.get(page_num + 1, [])

            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue

                block_html = []
                is_heading = False

                for line in block.get("lines", []):
                    line_html = []
                    for span in line.get("spans", []):
                        raw_text_span = span.get("text", "")
                        if not raw_text_span.strip():
                            if raw_text_span == " ":
                                line_html.append(" ")
                            continue

                        text = raw_text_span.replace("<", "&lt;").replace(">", "&gt;")

                        # Heading detection by font size
                        if span.get("size", 12.0) > avg_size + 1:
                            is_heading = True

                        # Bold / italic detection
                        flags = span.get("flags", 0)
                        font_name = span.get("font", "").lower()
                        is_bold = bool(flags & 16) or "bold" in font_name
                        is_italic = bool(flags & 2) or "italic" in font_name

                        # Link detection — check if span bbox overlaps with a hyperlink
                        span_bbox = span.get("bbox", (0, 0, 0, 0))
                        link_url = _span_in_link(span_bbox, page_links)

                        if link_url:
                            text = f'<a href="{link_url}" style="color:#6366f1;text-decoration:underline;">{text}</a>'
                        else:
                            if is_bold:
                                text = f"<strong>{text}</strong>"
                            if is_italic:
                                text = f"<em>{text}</em>"

                        line_html.append(text)

                    line_text = "".join(line_html).strip()

                    # ── Link-row reconstruction ─────────────────────────────
                    # When a PDF uses a single annotation rect for the whole
                    # link row, _span_in_link assigns every span to the same
                    # URL. Detect this by checking if 2+ pre-extracted links
                    # sit on this line (by y-overlap) and rebuild the HTML
                    # with individual <a> tags for each link.
                    if line_text and page_links and line.get("spans"):
                        all_bboxes = [
                            s.get("bbox", (0, 0, 0, 0))
                            for s in line.get("spans", [])
                            if s.get("text", "").strip()
                        ]
                        if all_bboxes:
                            ly0 = min(bb[1] for bb in all_bboxes)
                            ly1 = max(bb[3] for bb in all_bboxes)
                            margin = max((ly1 - ly0) * 0.6, 3.0)
                            row_links = [
                                lnk for lnk in page_links
                                if lnk.get("rect") and lnk.get("text") and lnk.get("url")
                                and lnk["rect"][1] <= ly1 + margin
                                and lnk["rect"][3] >= ly0 - margin
                            ]
                            if len(row_links) > 1:
                                row_links.sort(key=lambda l: l.get("position", 0))
                                parts: list = []
                                for i, lnk in enumerate(row_links):
                                    lnk_text = (lnk["text"]
                                                .replace("&", "&amp;")
                                                .replace("<", "&lt;")
                                                .replace(">", "&gt;"))
                                    sep = lnk.get("separator", "")
                                    parts.append(
                                        f'<a href="{lnk["url"]}" '
                                        f'style="color:#6366f1;text-decoration:underline;">'
                                        f'{lnk_text}</a>'
                                    )
                                    if sep and i < len(row_links) - 1:
                                        parts.append(
                                            sep.replace("&", "&amp;")
                                               .replace("<", "&lt;")
                                               .replace(">", "&gt;")
                                        )
                                line_text = "".join(parts)

                    if line_text:
                        block_html.append(line_text)

                if not block_html:
                    continue

                in_list = False
                prev_was_wrapped_line = False

                for line_text in block_html:
                    clean = re.sub(r'<[^>]+>', '', line_text).strip()
                    is_all_caps_bold = len(clean) < 80 and "<strong>" in line_text and clean.isupper()
                    is_title_like = (
                        len(clean) < 120 and (
                            "\u2014" in clean
                            or " \u2013 " in clean
                            or " — " in clean
                            or " – " in clean
                            or " · " in clean
                            or is_heading
                            or is_all_caps_bold
                        )
                    )

                    if line_text.startswith("&#x2022;") or line_text.startswith("•") or line_text.startswith("- "):
                        if not in_list:
                            html_parts.append("<ul>")
                            in_list = True
                        clean_text = line_text.replace("&#x2022;", "").replace("•", "").lstrip("- ").strip()
                        html_parts.append(f"<li>{clean_text}</li>")
                        prev_was_wrapped_line = False
                    else:
                        if in_list:
                            if is_title_like:
                                html_parts.append("</ul>")
                                in_list = False
                                if is_heading or is_all_caps_bold:
                                    html_parts.append(f"<h2>{line_text}</h2>")
                                else:
                                    html_parts.append(f"<h3>{line_text}</h3>")
                                prev_was_wrapped_line = False
                            else:
                                # Continuation of previous bullet
                                last_li = html_parts.pop()
                                html_parts.append(last_li[:-5] + " " + line_text + "</li>")
                        else:
                            if is_heading or is_all_caps_bold:
                                html_parts.append(f"<h2>{line_text}</h2>")
                                prev_was_wrapped_line = False
                            else:
                                html_parts.append(f"<p>{line_text}</p>")
                                prev_was_wrapped_line = True

                if in_list:
                    html_parts.append("</ul>")

        raw_html = "".join(html_parts).replace("</ul><ul>", "")

        # ── Join hyphenated line-breaks then clean artifacts in HTML text ───
        rich_text = _join_hyphens_in_html(raw_html)
        rich_text = _clean_html_text_nodes(rich_text)

    except Exception as e:
        logger.warning("PyMuPDF semantic HTML extraction failed: %s", e)

    # ── Raw text extraction (pdfplumber primary, PyMuPDF fallback) ──────────
    try:
        result = _parse_pdf_pdfplumber(file_path, links=links)
    except Exception as e:
        logger.warning("pdfplumber failed (%s), falling back to PyMuPDF", e)
        result = _parse_pdf_pymupdf(file_path, links=links)

    if not rich_text:
        rich_text = f"<p style='white-space: pre-wrap;'>{result['raw_text']}</p>"

    result["rich_text"] = rich_text
    return result


def _join_hyphens_in_html(html: str) -> str:
    """
    Apply hyphenated line-break joining only to text content between HTML tags,
    not to URLs or tag attributes.
    """
    # Split on tags — even indices are text, odd indices are tags
    parts = re.split(r'(<[^>]+>)', html)
    for i in range(0, len(parts), 2):  # only text nodes
        parts[i] = _join_hyphenated_linebreaks(parts[i])
    return "".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# pdfplumber extraction
# ─────────────────────────────────────────────────────────────────────────────


def _chars_to_spaced_text(chars: List[Dict]) -> str:
    """
    Convert pdfplumber character dicts to a plain text string, inserting
    a space wherever the horizontal gap between characters exceeds ~30%
    of the average character width.

    This is the Resume-Matcher approach: build text from character positions
    so that adjacent items on the same PDF line (e.g. an email address
    immediately followed by a hyperlink label) get a space inserted between
    them even when the PDF byte stream writes them without one.
    """
    if not chars:
        return ""

    # Sort by line then by x-position
    LINE_TOL = 4.5  # px — chars within this vertical band share a line
    # 3.0 was too tight for large-font section headers: characters in the same
    # visual line sometimes differ by up to 4 px in their 'top' coordinate when
    # the font is 14–18 pt, causing one line to be split into two ghost lines.

    # Group into lines by approximate y (top) value
    sorted_chars = sorted(chars, key=lambda c: (c.get("top", 0), c.get("x0", 0)))
    lines: List[List[Dict]] = []
    current: List[Dict] = []
    ref_top: float = -1.0

    for ch in sorted_chars:
        ch_top = ch.get("top", 0)
        if ref_top < 0 or abs(ch_top - ref_top) > LINE_TOL:
            if current:
                lines.append(current)
            current = [ch]
            ref_top = ch_top
        else:
            current.append(ch)
    if current:
        lines.append(current)

    text_lines: List[str] = []
    for line in lines:
        line = sorted(line, key=lambda c: c.get("x0", 0))
        buf: List[str] = []
        prev: Dict = {}
        for ch in line:
            char_text = ch.get("text", "")
            if not char_text or char_text in ("\x00", "\n", "\r"):
                continue
            if prev:
                gap = ch.get("x0", 0) - prev.get("x1", 0)
                # Average width of the two adjacent chars
                w_prev = prev.get("x1", 0) - prev.get("x0", 0)
                w_cur = ch.get("x1", 0) - ch.get("x0", 0)
                avg_w = (w_prev + w_cur) / 2.0 if (w_prev + w_cur) > 0 else 4.0
                # Insert a space if gap is more than 30% of avg char width
                # and we haven't already buffered a space
                if gap > avg_w * 0.30 and (not buf or buf[-1] != " "):
                    buf.append(" ")
            buf.append(char_text)
            prev = ch
        line_str = "".join(buf).strip()
        if line_str:
            text_lines.append(line_str)

    return "\n".join(text_lines)


def _parse_pdf_pdfplumber(file_path: str, links: List[Dict] = None) -> dict:
    import pdfplumber

    if links is None:
        links = _extract_pdf_links(file_path)

    pages_data = []
    all_text_parts = []
    char_data = []
    fonts_seen = set()
    has_tables = False
    has_header_footer = False
    has_images = False

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages):
            page_height = page.height or 1000

            # ── Primary: build text from character positions (Resume-Matcher
            #    approach) — properly inserts spaces at visual word gaps.
            # ── Fallback: extract_text(layout=True) / extract_text()
            page_chars = list(page.chars)
            if page_chars:
                try:
                    text = _chars_to_spaced_text(page_chars)
                except Exception:
                    text = ""
            else:
                text = ""
            if not text.strip():
                try:
                    text = page.extract_text(layout=True) or ""
                except TypeError:
                    text = page.extract_text() or ""

            all_text_parts.append(text)

            tables = page.extract_tables()
            if tables and len(tables) > 0 and any(t for t in tables):
                has_tables = True

            chars = page.chars
            for ch in chars:
                fn = ch.get("fontname", "")
                fs = ch.get("size", 0)
                if fn:
                    fonts_seen.add((fn, round(fs, 1)))
                top = ch.get("top", 0)
                bottom = ch.get("bottom", 0)
                if top < page_height * 0.08 or bottom > page_height * 0.92:
                    has_header_footer = True
                char_data.append({
                    "x": ch.get("x0", 0),
                    "y": ch.get("top", 0),
                    "fontname": fn,
                    "size": fs,
                    "text": ch.get("text", ""),
                })

            pages_data.append({"page": page_num + 1, "text": text, "char_count": len(chars)})

    raw_text = "\n".join(all_text_parts).strip()

    # ── Join hyphenated line-breaks, then run full normalization pipeline ────
    raw_text = _join_hyphenated_linebreaks(raw_text)
    raw_text = _normalize_raw_text(raw_text)

    # ── Append extracted link URLs to raw_text ───────────────────────────────
    # This ensures the ATS engine can see the actual URLs (portfolio, github, etc.)
    if links:
        seen_urls: set = set()
        link_lines = []
        for lnk in links:
            url = lnk.get("url", "").strip()
            label = lnk.get("text", "").strip()
            if url and url not in seen_urls:
                seen_urls.add(url)
                link_lines.append(f"{label}: {url}" if label else url)
        if link_lines:
            raw_text += "\n\n[EXTRACTED LINKS]\n" + "\n".join(link_lines)

    fonts_list = [{"fontname": fn, "size": fs} for fn, fs in fonts_seen]

    try:
        import fitz
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            if page.get_images(full=True):
                has_images = True
                break
        doc.close()
    except Exception:
        pass

    return {
        "raw_text": raw_text,
        "pages": pages_data,
        "page_count": page_count,
        "file_type": "pdf",
        "links": links,
        "metadata": {
            "fonts_used": [f["fontname"] for f in fonts_list],
            "font_count": len(fonts_seen),
            "font_details": fonts_list,
            "has_tables": has_tables,
            "has_images": has_images,
            "has_header_footer": has_header_footer,
            "has_special_characters": _has_special_chars(raw_text),
            "char_data": char_data[:500],
            "word_count": len(raw_text.split()),
            "extracted_links": links,
        },
    }


# ─────────────────────────────────────────────────────────────────────────────
# PyMuPDF fallback extraction
# ─────────────────────────────────────────────────────────────────────────────

def _parse_pdf_pymupdf(file_path: str, links: List[Dict] = None) -> dict:
    import fitz

    if links is None:
        links = _extract_pdf_links(file_path)

    pages_data = []
    all_text_parts = []
    has_images = False

    doc = fitz.open(file_path)
    page_count = len(doc)

    for page_num in range(page_count):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        all_text_parts.append(text)
        pages_data.append({"page": page_num + 1, "text": text, "char_count": len(text)})
        if page.get_images(full=True):
            has_images = True

    doc.close()

    raw_text = "\n".join(all_text_parts).strip()
    raw_text = _join_hyphenated_linebreaks(raw_text)
    raw_text = _normalize_raw_text(raw_text)

    if links:
        seen_urls: set = set()
        link_lines = []
        for lnk in links:
            url = lnk.get("url", "").strip()
            label = lnk.get("text", "").strip()
            if url and url not in seen_urls:
                seen_urls.add(url)
                link_lines.append(f"{label}: {url}" if label else url)
        if link_lines:
            raw_text += "\n\n[EXTRACTED LINKS]\n" + "\n".join(link_lines)

    return {
        "raw_text": raw_text,
        "pages": pages_data,
        "page_count": page_count,
        "file_type": "pdf",
        "links": links,
        "metadata": {
            "fonts_used": [],
            "font_count": 0,
            "font_details": [],
            "has_tables": False,
            "has_images": has_images,
            "has_header_footer": False,
            "has_special_characters": _has_special_chars(raw_text),
            "char_data": [],
            "word_count": len(raw_text.split()),
            "extracted_links": links,
        },
    }


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Parser
# ─────────────────────────────────────────────────────────────────────────────

def _parse_docx(file_path: str) -> dict:
    from docx import Document

    doc = Document(file_path)

    # Extract hyperlinks from docx relationships
    docx_links: List[Dict] = []
    try:
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                url = rel.target_ref
                if url and url.startswith("http"):
                    docx_links.append({"url": url, "text": ""})
    except Exception:
        pass

    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    raw_text = "\n".join(paragraphs)

    if docx_links:
        seen_urls: set = set()
        link_lines = []
        for lnk in docx_links:
            url = lnk.get("url", "").strip()
            if url and url not in seen_urls:
                seen_urls.add(url)
                link_lines.append(url)
        if link_lines:
            raw_text += "\n\n[EXTRACTED LINKS]\n" + "\n".join(link_lines)

    styles = set()
    for para in doc.paragraphs:
        if para.style and para.style.name:
            styles.add(para.style.name)

    has_header_footer = False
    try:
        for section in doc.sections:
            header_text = " ".join(p.text for p in section.header.paragraphs).strip()
            footer_text = " ".join(p.text for p in section.footer.paragraphs).strip()
            if header_text or footer_text:
                has_header_footer = True
                break
    except Exception:
        pass

    has_tables = len(doc.tables) > 0
    has_images = len(doc.inline_shapes) > 0

    has_text_boxes = False
    try:
        body_xml = doc.element.body.xml
        has_text_boxes = "txbxContent" in body_xml
    except Exception:
        pass

    # Generate rich HTML
    rich_text_parts = []
    _in_list = False
    for para in doc.paragraphs:
        if para.text.strip():
            style_name = para.style.name if para.style else ""
            is_bullet = style_name in ('List Bullet', 'List Bullet 2', 'List Bullet 3',
                                       'List Paragraph') or style_name.startswith('List')
            if "Heading 1" in style_name:
                if _in_list:
                    rich_text_parts.append("</ul>")
                    _in_list = False
                rich_text_parts.append(f"<h2>{para.text}</h2>")
            elif "Heading" in style_name:
                if _in_list:
                    rich_text_parts.append("</ul>")
                    _in_list = False
                rich_text_parts.append(f"<h3>{para.text}</h3>")
            else:
                html_para = []
                for run in para.runs:
                    run_text = run.text.replace("<", "&lt;").replace(">", "&gt;")
                    if run.bold:
                        run_text = f"<strong>{run_text}</strong>"
                    if run.italic:
                        run_text = f"<em>{run_text}</em>"
                    html_para.append(run_text)
                if is_bullet:
                    # Open <ul> only once; consecutive bullets share the same list
                    if not _in_list:
                        rich_text_parts.append("<ul>")
                        _in_list = True
                    rich_text_parts.append(f"<li>{''.join(html_para)}</li>")
                else:
                    if _in_list:
                        rich_text_parts.append("</ul>")
                        _in_list = False
                    rich_text_parts.append(f"<p>{''.join(html_para)}</p>")
        else:
            # Empty paragraph — close any open list
            if _in_list:
                rich_text_parts.append("</ul>")
                _in_list = False
    if _in_list:
        rich_text_parts.append("</ul>")

    for table in doc.tables:
        rich_text_parts.append("<table border='1'>")
        for row in table.rows:
            rich_text_parts.append("<tr>")
            for cell in row.cells:
                if cell.text.strip():
                    rich_text_parts.append(f"<td>{cell.text}</td>")
            rich_text_parts.append("</tr>")
        rich_text_parts.append("</table>")

    return {
        "raw_text": raw_text,
        "rich_text": "".join(rich_text_parts),
        "pages": [{"page": 1, "text": raw_text, "char_count": len(raw_text)}],
        "page_count": 1,
        "file_type": "docx",
        "links": docx_links,
        "metadata": {
            "fonts_used": list(styles),
            "font_count": len(styles),
            "font_details": [{"fontname": s, "size": 0} for s in styles],
            "has_tables": has_tables or has_text_boxes,
            "has_images": has_images,
            "has_header_footer": has_header_footer,
            "has_special_characters": _has_special_chars(raw_text),
            "char_data": [],
            "word_count": len(raw_text.split()),
            "has_text_boxes": has_text_boxes,
            "extracted_links": docx_links,
        },
    }


# ─────────────────────────────────────────────────────────────────────────────
# TXT Parser
# ─────────────────────────────────────────────────────────────────────────────

def _parse_txt(file_path: str) -> dict:
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        raw_text = f.read()

    # Extract any bare URLs from the text
    url_pattern = re.compile(r'https?://[^\s\)\]\>\"\']+')
    txt_links = [{"url": u, "text": ""} for u in url_pattern.findall(raw_text)]

    rich_text_parts = []
    for line in raw_text.split("\n"):
        stripped = line.strip()
        if stripped:
            if stripped.isupper() and len(stripped) < 50:
                rich_text_parts.append(f"<h2>{stripped}</h2>")
            elif stripped.startswith("- ") or stripped.startswith("* "):
                rich_text_parts.append(f"<ul><li>{stripped[2:]}</li></ul>")
            else:
                # Linkify bare URLs
                line_html = url_pattern.sub(
                    lambda m: f'<a href="{m.group()}" style="color:#6366f1;">{m.group()}</a>',
                    stripped.replace("<", "&lt;").replace(">", "&gt;"),
                )
                rich_text_parts.append(f"<p>{line_html}</p>")

    return {
        "raw_text": raw_text,
        "rich_text": "".join(rich_text_parts),
        "pages": [{"page": 1, "text": raw_text, "char_count": len(raw_text)}],
        "page_count": 1,
        "file_type": "txt",
        "links": txt_links,
        "metadata": {
            "fonts_used": [],
            "font_count": 0,
            "font_details": [],
            "has_tables": False,
            "has_images": False,
            "has_header_footer": False,
            "has_special_characters": _has_special_chars(raw_text),
            "char_data": [],
            "word_count": len(raw_text.split()),
            "extracted_links": txt_links,
        },
    }


# ─────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────────────────────────────────────

def _has_special_chars(text: str) -> bool:
    for ch in text:
        if ch == '\x00' or (ord(ch) < 32 and ch not in '\n\r\t'):
            return True
    return False
